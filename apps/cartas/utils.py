#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os.path
import PyPDF2
import tabula
import numpy
import csv
from decimal import Decimal
from datetime import date
from django.utils import timezone
from docx import Document
from .models import Plazo

BASE = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'media')

# YYYY/MM/DD
def calcularEdad(fecha): 
    today = date.today() 
    age = today.year - fecha.year - ((today.month, today.day) < (fecha.month, fecha.day)) 
    return age 

def get_meses_plazo(id_plazo):
    meses_plazo = 0

    if id_plazo == 1:
        meses_plazo = 24
    elif id_plazo == 2:
        meses_plazo = 36
    elif id_plazo == 4:
        meses_plazo = 60
    else:
        meses_plazo = 48
    return meses_plazo

def get_id_meses_plazo(meses_plazo):

    if meses_plazo == 24:
        id_plazo = 0
    elif meses_plazo == 36:
        id_plazo = 1
    elif meses_plazo == 60:
        id_plazo = 3
    else:
        id_plazo = 2
    return id_plazo

class ObtenerDatosPDFMixin(object):
    contenido = []
    montos = []

    def convertir_csv(self, ruta_pdf):
            df = tabula.read_pdf(ruta_pdf ,multiple_tables=True)
            # Convertimos el archivo a formato CSV y lo guardamos 
            ruta_csv = os.path.join(BASE, str("reporte_pensionado_actual.csv"))
            ruta_csv = ruta_csv.replace('/apps','')
            tabula.convert_into(ruta_pdf, ruta_csv, output_format="csv", stream=True, pages=1)

            return ruta_csv

    def leer_csv(self, ruta_csv):
        self.contenido = []
        f = open(ruta_csv)
        
        csv_f = csv.reader(f)
        for index, row in enumerate(csv_f):
            self.contenido.append(list(row[i] for i in [0,1,2,3,4]))

    def obtener_datos_csv(self):
        contenido_array = numpy.array(self.contenido)

        nss = numpy.where(contenido_array == 'NSS')
        numero_social = self.contenido[nss[0][0] - 1][0]

        nombre_titulo = numpy.where(contenido_array == 'NOMBRE DEL ASEGURADO')
        nombre = self.contenido[nombre_titulo[0][0] - 1][2]

        try:
            estado_titulo = numpy.where(contenido_array == 'DELEGACI?N DE PAGO')
            estado = self.contenido[estado_titulo[0][0] - 1][3]
        except:
            estado_titulo = numpy.where(contenido_array == 'FECHA DE  VENCIMIENTO DELEGACI?N DE PAGO')
            estado = self.contenido[estado_titulo[0][0] - 1][2][16:]

        ciudad_titulo = numpy.where(contenido_array == 'SUBDELEGACI?N DE PAGO')
        ciudad = self.contenido[ciudad_titulo[0][0] - 1][4]

        domicilio_titulo = numpy.where(contenido_array == 'DOMICILIO')
        domicilio = self.contenido[domicilio_titulo[0][0] - 1][2]
        if domicilio == '':
            domicilio = self.contenido[domicilio_titulo[0][0] - 1][1]
            if domicilio == '':
                domicilio = self.contenido[domicilio_titulo[0][0] - 1][0]

        curp_titulo = numpy.where(contenido_array == 'CURP')
        curp = self.contenido[curp_titulo[0][0] - 1][4]
        if curp == '':
            curp = self.contenido[curp_titulo[0][0] - 1][3]
        fecha_nacimiento_curp = curp[4:10]
        # Se espera que todos los registros contengan una fecha de nacimiento a mayor al año 2000 (en el 2020)
        # el CURP solo devuelve 2 dígitos del año de nacimiento, agregamos manualmente el 19** faltante
        # print(fecha_nacimiento_curp)
        # print(fecha_nacimiento_curp[0:2])
        # print(fecha_nacimiento_curp[2:4])
        # print(fecha_nacimiento_curp[4:])

        try:
            fecha_nacimiento = date(int('19' + fecha_nacimiento_curp[0:2]), int(fecha_nacimiento_curp[2:4]), int(fecha_nacimiento_curp[4:]))
        except:
            fecha_nacimiento = date(int('19' + fecha_nacimiento_curp[0:2]), int(fecha_nacimiento_curp[2:4]), int(fecha_nacimiento_curp[4:]))
        edad = calcularEdad(fecha_nacimiento)

        liquido_titulo = numpy.where(contenido_array == 'L?QUIDO')
        liquido = self.contenido[liquido_titulo[0][0] - 1][4][2:]

        try:
            capacidad_titulo = numpy.where(numpy.char.find(contenido_array, 'CAPACIDAD DE CR?DITO')>=0)
            capacidad = self.contenido[capacidad_titulo[0][0] - 1][2][2:]
            capacidad = capacidad.split('$')[0].strip()
        except: 
            capacidad = '-9999.0'

        deuda_concepto = numpy.where(numpy.char.find(contenido_array, '301 PRESTAMO')>=0)

        if deuda_concepto[0].any():
            deuda_empresa = self.contenido[deuda_concepto[0][-1]][0][15:]
            deuda_cantidad_pagos = self.contenido[deuda_concepto[0][-1]][3]
            deuda_cantidad = self.contenido[deuda_concepto[0][-1]][4][3:]

            self.datos_pensionado['deuda_empresa'] = deuda_empresa
            self.datos_pensionado['deuda_cantidad_pagos'] = deuda_cantidad_pagos
            self.datos_pensionado['deuda_cantidad'] = deuda_cantidad

        self.datos_pensionado['numero_social'] = numero_social
        self.datos_pensionado['nombre'] = nombre
        self.datos_pensionado['estado'] = estado
        self.datos_pensionado['ciudad'] = ciudad
        self.datos_pensionado['domicilio'] = domicilio
        self.datos_pensionado['edad'] = edad
        self.datos_pensionado['liquido'] = liquido
        self.datos_pensionado['capacidad'] = capacidad


    def sobreescribir_carta(self, nombre_pensionado, cantidad_prestamo):
        # ruta_carta = '/Users/macbookair/Reportes/reportes/static/media/formatos/formato_carta.docx'
        # BASE = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'media')
        ruta_carta = os.path.join(BASE, str("formatos/formato_carta.docx"))
        ruta_carta = ruta_carta.replace('/apps','')

        doc = Document(ruta_carta)
        for p in doc.paragraphs:
            if 'NOMBRE' in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if 'NOMBRE' in inline[i].text:
                        text = inline[i].text.replace('NOMBRE', nombre_pensionado)
                        inline[i].text = text
            if 'CANTIDAD' in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if 'CANTIDAD' in inline[i].text:
                        text = inline[i].text.replace('CANTIDAD', cantidad_prestamo)
                        inline[i].text = text
        today = date.today() 
        fecha = str(today.year) + str(today.month)    
        # desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') 
        desktop = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop') 

        ruta_guardado = os.path.join(desktop, str("cartas/"+ fecha))
        ruta_guardado = ruta_guardado.replace('/apps','')
        
        # ruta_guardado = '/Users/macbookair/Reportes/reportes/static/media/cartas/'+ fecha
        
        if not os.path.exists(ruta_guardado):
            os.makedirs(ruta_guardado)
        doc.save(ruta_guardado + '/carta_' + nombre_pensionado + '.docx')

    def sobreescribir_sobre(self, nombre_pensionado, direccion_envio):
        # ruta_sobre = '/Users/macbookair/Reportes/reportes/static/media/formatos/formato_sobre.docx'
        # BASE = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'media')
        ruta_sobre = os.path.join(BASE, str("formatos/formato_sobre.docx"))
        ruta_sobre = ruta_sobre.replace('/apps','')

        direccion_envio = direccion_envio.replace(', ', ',')
        direccion_envio = direccion_envio.replace(',', ',\n')

        doc = Document(ruta_sobre)
        for p in doc.paragraphs:
            if 'NOMBRE' in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if 'NOMBRE' in inline[i].text:
                        text = inline[i].text.replace('NOMBRE', nombre_pensionado)
                        inline[i].text = text
            if 'DIRECCION' in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if 'DIRECCION' in inline[i].text:
                        text = inline[i].text.replace('DIRECCION', direccion_envio)
                        inline[i].text = text
        today = date.today() 
        fecha = str(today.year) + str(today.month)

        # desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') 
        desktop = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop') 

        ruta_guardado = os.path.join(desktop, str("sobres/"+ fecha))
        ruta_guardado = ruta_guardado.replace('/apps','')
        
        # ruta_guardado = '/Users/macbookair/Reportes/reportes/static/media/sobres/'+ fecha
        
        if not os.path.exists(ruta_guardado):
            os.makedirs(ruta_guardado)
        
        doc.save(ruta_guardado + '/sobre_' + nombre_pensionado + '.docx')
        

    def generar_archivo_cotizador(self, ruta_pdf):
        ruta_csv = os.path.join(BASE, str("cotizadores/cotizador-41.0.csv"))
        ruta_csv = ruta_csv.replace('/apps','')
        tabula.convert_into(ruta_pdf, ruta_csv, output_format="csv", stream=True, pages='all')

        return ruta_csv

    def modificar_cotizador_csv(self, ruta_csv):
        listado_textos = ['MESES', 'Tasa Nominal Anual', 'CAT', 'Monto Solicitado']
        f = open(ruta_csv)
        
        csv_f = csv.reader(f)
        for index, row in enumerate(csv_f):
            self.contenido.append(list(row[i] for i in [0,1,2,3,4]))
        for row in self.contenido:
            if row[0] not in listado_textos:
                self.montos.append(row)
        
        with open(ruta_csv, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(self.montos)
   
    def generar_plazos_actuales(self, ruta_csv):
        # Actualizamos los plazos actuales en la base de datos a 99 
        # para no mostrarlos en el resultado cuando se consulten
        Plazo.objects.all().update(estatus_registro=99, actualizado=timezone.now())
        f = open(ruta_csv)
        csv_f = csv.reader(f)
        for row in csv_f:
            monto = float(row[0][1:].replace(",",""))

            for i, col in enumerate(row):
                if i == 1:
                    plazo = Plazo.objects.create()
                    plazo.meses_plazo = 24
                    plazo.monto_solicitado = monto
                    plazo.pago = float(col[1:].replace(",",""))
                    plazo.save()
                
                if i == 2:
                    plazo = Plazo.objects.create()
                    plazo.meses_plazo = 36
                    plazo.monto_solicitado = monto
                    plazo.pago = float(col[1:].replace(",",""))
                    plazo.save()

                if i == 3:
                    plazo = Plazo.objects.create()
                    plazo.meses_plazo = 48
                    plazo.monto_solicitado = monto
                    plazo.pago = float(col[1:].replace(",",""))
                    plazo.save()

                if i == 4:
                    plazo = Plazo.objects.create()
                    plazo.meses_plazo = 60
                    plazo.monto_solicitado = monto
                    plazo.pago = float(col[1:].replace(",",""))
                    plazo.save()
